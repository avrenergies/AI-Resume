<<<<<<< HEAD
from docx import Document


def extract_paragraphs(doc):
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    return texts


def extract_tables(doc):
    texts = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                texts.append(" | ".join(row_text))

    return texts


def extract_headers_footers(doc):
    texts = []

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        for para in footer.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

    return texts


def clean_text(text):
    # Remove excessive blank lines
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    return "\n".join(lines)


def docx_to_text(file_path):

    try:
        doc = Document(file_path)
    except Exception:
        return ""

    content = []

    # 1️⃣ Paragraphs
    content.extend(extract_paragraphs(doc))

    # 2️⃣ Tables (VERY IMPORTANT FOR RESUMES)
    content.extend(extract_tables(doc))

    # 3️⃣ Header & Footer
    content.extend(extract_headers_footers(doc))

    final_text = "\n".join(content)

    return clean_text(final_text)
=======
from docx import Document


def extract_paragraphs(doc):
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    return texts


def extract_tables(doc):
    texts = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                texts.append(" | ".join(row_text))

    return texts


def extract_headers_footers(doc):
    texts = []

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for para in header.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        for para in footer.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

    return texts


def clean_text(text):
    # Remove excessive blank lines
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    return "\n".join(lines)


def docx_to_text(file_path):

    try:
        doc = Document(file_path)
    except Exception:
        return ""

    content = []

    # 1️⃣ Paragraphs
    content.extend(extract_paragraphs(doc))

    # 2️⃣ Tables (VERY IMPORTANT FOR RESUMES)
    content.extend(extract_tables(doc))

    # 3️⃣ Header & Footer
    content.extend(extract_headers_footers(doc))

    final_text = "\n".join(content)

    return clean_text(final_text)
>>>>>>> 9d2d7face242eebb6a4a3d878a35ead4285cf42d
